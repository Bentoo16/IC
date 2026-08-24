import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from io import BytesIO
import re

# ---------------------------------------------------------------------------
# CSS customizado
# ---------------------------------------------------------------------------
st.markdown("""
<style>
    .tabela-respostas {
        width: 100%;
        border-collapse: collapse;
        font-size: 0.85rem;
        margin-bottom: 1.5rem;
    }
    .tabela-respostas th {
        background: #E0E0E0;
        color: #333;
        padding: 0.5rem;
        text-align: center;
        border: 1px solid #ccc;
    }
    .tabela-respostas th.caso-header {
        color: #FF0000;
    }
    .tabela-respostas td {
        padding: 0.4rem 0.5rem;
        border: 1px solid #dee2e6;
        text-align: center;
    }
    .tabela-respostas td:first-child {
        text-align: left;
        font-weight: 500;
    }
    .preview-box {
        background: #f4f4f4;
        border-left: 4px solid #555;
        padding: 1rem 1.2rem;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    .progresso-texto {
        font-size: 0.85rem;
        color: #555;
        margin-bottom: 0.3rem;
    }
    .stButton>button {
        font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Funções auxiliares
# ---------------------------------------------------------------------------
def extrair_numero(nome):
    match = re.search(r'\d+', nome)
    return int(match.group()) if match else 0


def gerar_tabela_html(casos_ordenados, perguntas_ordenadas, escolhas_casos):
    html = "<table class='tabela-respostas'>"
    html += "<tr><th rowspan='2'>Pergunta</th>"
    for caso in casos_ordenados:
        html += f"<th class='caso-header' colspan='2'>{caso}</th>"
    html += "</tr><tr>"
    for _ in casos_ordenados:
        html += "<th>Sim</th><th>Não</th>"
    html += "</tr>"
    for pergunta in perguntas_ordenadas:
        html += "<tr>"
        html += f"<td>{pergunta}</td>"
        for caso in casos_ordenados:
            item = escolhas_casos.get(caso, {}).get(pergunta, {})
            resposta = item.get("resposta", "-") if isinstance(item, dict) else item
            sim_cell = "X" if resposta == "Sim" else ""
            nao_cell = "X" if resposta == "Não" else ""
            html += f"<td>{sim_cell}</td><td>{nao_cell}</td>"
        html += "</tr>"
    html += "</table>"
    return html


def contar_perguntas(grupos):
    total = 0
    for qs in grupos.values():
        total += len(qs)
    return total


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# Configuração da IA
# ---------------------------------------------------------------------------
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
model = genai.GenerativeModel('gemini-3.1-flash-lite')

st.title("Gerador de Relatórios - Mamografia")

# ---------------------------------------------------------------------------
# Inicialização do session_state
# ---------------------------------------------------------------------------
if "dados_cabecalho" not in st.session_state:
    st.session_state.dados_cabecalho = {
        "mamografo_fabricante": "",
        "mamografo_modelo": "",
        "cnes": "",
        "qiid": "",
        "tipo_mamografo": None,
        "instituicao": "",
        "cidade": "",
        "estado": "",
        "servico": "",
    }
if "servico" not in st.session_state.dados_cabecalho:
    st.session_state.dados_cabecalho["servico"] = ""
if "casos_salvos" not in st.session_state:
    st.session_state.casos_salvos = {}
if "relatorios_ia" not in st.session_state:
    st.session_state.relatorios_ia = {}
if "relatorio_geral_salvo" not in st.session_state:
    st.session_state.relatorio_geral_salvo = None
if "consideracoes_caso" not in st.session_state:
    st.session_state.consideracoes_caso = {}
if "consideracoes_gerais" not in st.session_state:
    st.session_state.consideracoes_gerais = ""
if "escolhas_casos" not in st.session_state:
    st.session_state.escolhas_casos = {}
if "identificacao_exames" not in st.session_state:
    st.session_state.identificacao_exames = {}

# ---------------------------------------------------------------------------
# Barra de progresso da sessão
# ---------------------------------------------------------------------------
total_casos = 5
casos_feitos = len(st.session_state.casos_salvos)
if casos_feitos > 0:
    st.markdown(f"<div class='progresso-texto'>Progresso: {casos_feitos} de {total_casos} casos analisados</div>", unsafe_allow_html=True)
    st.progress(casos_feitos / total_casos)

# ---------------------------------------------------------------------------
# Cabeçalho - dados da instituição
# ---------------------------------------------------------------------------
st.markdown("---")
st.subheader("Dados do Cabeçalho")

col1, col2 = st.columns(2)
with col1:
    fabricante = st.text_input("Mamógrafo - Fabricante:", value=st.session_state.dados_cabecalho["mamografo_fabricante"])
with col2:
    modelo = st.text_input("Mamógrafo - Modelo:", value=st.session_state.dados_cabecalho["mamografo_modelo"])

cnes = st.text_input("CNES:", value=st.session_state.dados_cabecalho["cnes"])
qiid = st.text_input("QIID:", value=st.session_state.dados_cabecalho["qiid"])

tipo = st.radio(
    "Tipo de mamógrafo:",
    ["Convencional", "Digital CR", "Digital DR", "DR retrofit"],
    index=0
    if st.session_state.dados_cabecalho["tipo_mamografo"] is None
    else ["Convencional", "Digital CR", "Digital DR", "DR retrofit"].index(
        st.session_state.dados_cabecalho["tipo_mamografo"]
    ),
    horizontal=True,
    key="tipo_mamografo_radio",
)

instituicao = st.text_input("Instituição:", value=st.session_state.dados_cabecalho["instituicao"])

col3, col4 = st.columns(2)
with col3:
    cidade = st.text_input("Cidade:", value=st.session_state.dados_cabecalho["cidade"])
with col4:
    estado = st.text_input("Estado:", value=st.session_state.dados_cabecalho["estado"])

servico = st.text_input("Serviço:", value=st.session_state.dados_cabecalho["servico"])

st.session_state.dados_cabecalho = {
    "mamografo_fabricante": fabricante,
    "mamografo_modelo": modelo,
    "cnes": cnes,
    "qiid": qiid,
    "tipo_mamografo": tipo,
    "instituicao": instituicao,
    "cidade": cidade,
    "estado": estado,
    "servico": servico,
}

# ---------------------------------------------------------------------------
# Biblioteca de perguntas (organizada por grupos)
# ---------------------------------------------------------------------------
perguntas = {
    "Avaliação dos Critérios de Posicionamento": {
        "Identificação correta do exame": {
            "opcoes": {
                "Sim": "",
                "Não": "A identificação das imagens enviadas para avaliação não está correta porque há texto impresso sobre áreas das mamas."
            },
        },
        "Adequada compressão de mama": {
            "opcoes": {
                "Sim": " ",
                "Não": "As imagens da mama deste caso estão com acentuada perda de definição das estruturas anatômicas (imagens tremidas) possivelmente causada pela pouca compressão da mama ou por movimentação da paciente durante a aquisição das imagens. "
            },
        },
        "Mamilo paralelo ao filme": {
            "opcoes": {
                "Sim": " ",
                "Não": "Adicionalmente, nestas incidências, os mamilos não estão perfilados paralelos ao filme. "
            },
        },
        "Visibilização completa do parênquima mamário": {
            "opcoes": {
                "Sim": " ",
                "Não": "Este mal posicionamento das mamas da paciente nas incidências não fornece uma visibilização completa do parênquima mamário, podendo prejudicar o diagnóstico de lesões em tecidos mamários de interesse. "
            },
        },
        "Músculo grande peitoral na altura do mamilo ou abaixo - na 0ML": {
            "opcoes": {
                "Sim": " ",
                "Não": "As imagens das incidências mediolaterais oblíquas (MLO) deste caso não incluem o músculo grande peitoral na altura do mamilo ou abaixo e as mamas para as incidências craniocaudais (CC) não estão bem posicionadas. Este mal posicionamento das mamas da paciente pode prejudicar o diagnóstico devido à visibilização incompleta de tecidos mamários de interesse. "
            },
        },
        "Prega inframamária incluída na radiografia - na 0ML": {
            "opcoes": {
                "Sim": " ",
                "Não": "As imagens das incidências mediolaterais oblíquas (MLO) não incluem a prega inframamária. Por isso as imagens das incidências MLO não mostram a visibilização completa do mamário."
            },
        },
    },
    "Avaliação dos Critérios Clínicos de Qualidade da Imagem": {
        "Visibilização adequada da pele (ausência na convencional ou presença na digital)":{
            "opcoes": {
                "Sim": " ",
                "Não": " "
            },
            "sub_opcoes": {
                "Dobra de pele junto a parede toráxica e papila não perfilada": "Na parte inferior da imagem desta incidência da mama (MLO) se observa uma dobra de pele junto à parede torácica e na imagem da mama (MLO) a papila não está perfilada em relação ao detector de imagem.",
                "Assimetria difusa da mama associada a espessamento da pele e do complexo areolopapilar.": "Neste caso, há assimetria difusa da mama associada a espessamento da pele e do complexo areolopapilar."
            },
        },
        "Visibilização das estruturas vasculares através do parênquima denso": {
            "opcoes": {
                "Sim": " ",
                "Não": " "
             },
        },
        "Visibilização dos ligamentos de Cooper": {
            "opcoes": {
                "Sim": " ",
                "Não": " "
            },
        },
        "As microcalcificações representa lesão verdadeira? (se houver lesão)": {
            "opcoes": {
                "Sim": " ",
                "Não": " "
            },
        },
        "A opacidade representa lesão verdadeira? (se houver lesão)": {
            "opcoes": {
                "Sim": " ",
                "Não": " "
            },
        },
        "O tecido glandular está adequadamente claro": {
            "opcoes": {
                "Sim": " ",
                "Não": " "
            },
        },
    },
    "Avaliação dos Critérios de Laudos": {
        "Resumo da história presente": {
            "opcoes": {"Sim": " ", "Não": "É importante que nos laudos conste a indicação do exame. Essa indicação deve conter uma história resumida da paciente (exame de rastreamento x diagnóstico / história familiar / antecedentes cirúrgicos e resultados de biópsias / sintomas e queixas da paciente ... )."},
        },
        "Utiliza corretamente o Léxico BI-RADS ou SISMAMA": {
            "opcoes": {"Sim": " ", "Não": "No laudo deste exame não foi utilizado corretamente o léxico do BI-RADS® ou do SISMAMA."},
        },
        "Classifica corretamente o exame segundo o BI-RADS": {
            "opcoes": {"Sim": " ", "Não": "O exame não foi classificado corretamente."},
        },
        "Recomendação correta segundo o BI-RADS": {
            "opcoes": {"Sim": "Recomenda corretamente o exame segundo o BI-RADS.", "Não": "No laudo deste exame não consta a recomendação de conduta em relação ao achado radiográfico reportado."},
        },
        "Interpretou corretamente todos os achados do exame": {
            "opcoes": {"Sim": "Interpretou corretamente todos os achados do exame.", "Não": "Não interpretou corretamente todos os achados do exame."},
        },
    },
    "Aspectos Físicos da Imagem": {
        "Contraste adequado": {
            "opcoes": {"Sim": " ", "Não": " "},
            "sub_opcoes": {
                "Contraste alto": (
                    "As imagens estão com o contraste aumentado devido à acentuada diferença "
                    "entre os tons de cinza claros e escuros presentes. Para que o contraste "
                    "das imagens seja considerado adequado, essa diferença deve ser menos acentuada."
                ),
                "Contraste muito alto": (
                    "As imagens deste exame estão com o contraste muito alto e com as regiões "
                    "correspondentes a tecidos mamários mais densos com os tons de cinza claro "
                    "saturados (muito claros, quase transparentes), o mesmo ocorrendo nas regiões "
                    "das axilas nas incidências mediolaterais oblíquas (MLO). Este aspecto das "
                    "imagens dificulta ou mesmo inviabiliza a identificação de microcalcificações "
                    "nas regiões de tecidos densos. Portanto, as imagens impressas enviadas para "
                    "avaliação foram consideradas sem qualidade técnica para a interpretação diagnóstica."
                ),
                "Contraste baixo": (
                    "As imagens das quatro incidências estão muito claras e, por conseguinte, com "
                    "o contraste reduzido devido à pouca diferença entre os tons de cinza claros "
                    "(regiões de tecido fibroglandular) e de cinza escuros (regiões de tecido "
                    "subcutâneo e tecido adiposo retromamário) presentes. Para que o contraste das "
                    "imagens seja considerado adequado, essa diferença deve ser mais acentuada."
                ),
            },
        },
        "Definição de estruturas": {
            "opcoes": {
                "Sim": "As estruturas estão bem definidas na imagem.",
                "Não": "As imagens da mama deste caso estão com acentuada perda de definição das estruturas anatômicas (imagens tremidas) possivelmente causada pela pouca compressão da mama ou por movimentação da paciente durante a aquisição das imagens.",
            },
        },
        "Saturação correta nas áreas claras": {
            "opcoes": {
                "Sim": "A imagem está bem saturada nas áreas claras.",
                "Não": "A imagem não está bem saturada nas áreas claras.",
            },
        },
        "Saturação correta nas áreas escuras": {
            "opcoes": {
                "Sim": "A imagem está bem saturada nas áreas escuras.",
                "Não": "A imagem não está bem saturada nas áreas escuras.",
            },
        },
        "Imagem sem ruído": {
            "opcoes": {"Sim": " ", "Não": " "},
        },
        "A área de fundo está adequadamente escura (enegrecimento película)": {
            "opcoes": {
                "Sim": " ",
                "Não": " ",
            },
        },
        "Imagem sem artefatos (se houver, descrever)": {
            "opcoes": {"Sim": " ", "Não": " "},
            #"gatilho_sub_opcoes": "Sim",
            "sub_opcoes": {
                "Possui artefatos na forma de linhas finas gerados pelo movimento insuficiente da grade antidifusora do mamógrafo. ": "Adicionalmente, elas apresentam diversos artefatos na forma de finas linhas verticais de tons de cinza claro gerados pelo movimento insuficiente da grade antidifusora do mamógrafo. Estas linhas causam uma impressão de ruído (aspecto granulado) perceptível nas imagens das pacientes.",
                "Possui artefatos decorrentes de desgastes e/ou danificadas. ": "As imagens enviadas para avaliação apresentam inúmeros artefatos de diversos tipos decorrentes das placas de imagem (IP) desgastadas e/ou danificadas.",
                "Possui escala métrica sobre as imagens das mamas. ": "Por fim, as imagens da mama têm uma escala métrica impressa na lateral do filme próxima à parede torácica da paciente. Estas escalas métricas impressas sobre as imagens das mamas constituem artefatos que devem ser retirados. "
            },
        },
    },
}

# ---------------------------------------------------------------------------
# Seleção do caso e perguntas
# ---------------------------------------------------------------------------
st.markdown("---")
caso_atual = st.selectbox("Escolha o Caso que vai analisar agora:", [1, 2, 3, 4, 5])
nome_caso = f"Caso {caso_atual}"

respostas_temporarias = []
abas_grupos = st.tabs(list(perguntas.keys()))
for idx, (nome_grupo, questoes) in enumerate(perguntas.items()):
    with abas_grupos[idx]:
        for titulo, info in questoes.items():
            st.subheader(titulo)
            escolha = st.radio("Selecione:", list(info["opcoes"].keys()), key=f"radio_{titulo}_c{caso_atual}", horizontal=True)
            gatilho = info.get("gatilho_sub_opcoes", "Não")
            sub_escolha = []
            if "sub_opcoes" in info and escolha == gatilho:
                sub_escolha = st.multiselect("Especifique:", list(info["sub_opcoes"].keys()), key=f"sub_{titulo}_c{caso_atual}")
            obs = st.text_input("Considerações adicionais:", key=f"obs_{titulo}_c{caso_atual}", placeholder="Opcional")
            respostas_temporarias.append({"titulo": titulo, "escolha": escolha, "sub_escolha": sub_escolha, "obs": obs})

st.markdown("---")
id_exame = st.text_input(
    "Identificação do Exame:",
    value=st.session_state.identificacao_exames.get(nome_caso, ""),
    key=f"id_exame_c{caso_atual}",
)
consideracoes_caso = st.text_area(
    "Considerações adicionais para este caso (opcional):",
    value=st.session_state.consideracoes_caso.get(nome_caso, ""),
    key=f"consideracoes_c{caso_atual}",
    height=100,
)

caso_ja_existe = nome_caso in st.session_state.casos_salvos
confirmacao = True
if caso_ja_existe:
    st.warning(f"O {nome_caso} já foi salvo anteriormente.")
    confirmacao = st.checkbox("Deseja sobrescrever o relatório existente?", key=f"conf_{caso_atual}")

if st.button(f"Analisar e Salvar {nome_caso}", type="primary", use_container_width=True):
    if caso_ja_existe and not confirmacao:
        st.warning("Marque a confirmação para sobrescrever o caso.")
    else:
        respostas_finais = []
        for item in respostas_temporarias:
            # Encontrar o grupo e a pergunta
            for questoes in perguntas.values():
                if item["titulo"] in questoes:
                    info_pergunta = questoes[item["titulo"]]
                    break
            gatilho_pergunta = info_pergunta.get("gatilho_sub_opcoes", "Não")
            if item["escolha"] == gatilho_pergunta and item["sub_escolha"]:
                frase_base = " ".join(
                    info_pergunta["sub_opcoes"][opcao] for opcao in item["sub_escolha"]
                )
            else:
                frase_base = info_pergunta["opcoes"][item["escolha"]]
            if item["obs"]:
                frase_base += f" Detalhe adicional: {item['obs']}"
            respostas_finais.append(frase_base)

        texto_bruto = " ".join(respostas_finais)
        texto_para_ia = texto_bruto
        if consideracoes_caso.strip():
            texto_para_ia += f"\n\n{consideracoes_caso}"

        escolhas = {
            item["titulo"]: {
                "resposta": item["escolha"],
                "sub_opcao": item["sub_escolha"],
            }
            for item in respostas_temporarias
        }

        with st.spinner("IA está formatando o relatório..."):
            try:
                prompt = (
                    f"Deixe essas frases em um único texto coeso, não é necessário acrescentar nada, apenas o texto coeso é o suficiente. "
                    f"Não mude as frases, apenas deixe o texto coeso para o {nome_caso}: {texto_para_ia}"
                )
                response = model.generate_content(prompt)

                st.session_state.casos_salvos[nome_caso] = texto_bruto
                st.session_state.consideracoes_caso[nome_caso] = consideracoes_caso
                st.session_state.identificacao_exames[nome_caso] = id_exame
                st.session_state.escolhas_casos[nome_caso] = escolhas
                st.session_state.relatorios_ia[nome_caso] = response.text
                st.session_state.docx_bytes = None

                st.success(f"{nome_caso} processado com sucesso!")
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao gerar relatório: {e}")

# ---------------------------------------------------------------------------
# Considerações gerais
# ---------------------------------------------------------------------------
st.markdown("---")
st.subheader("Considerações Gerais")
st.session_state.consideracoes_gerais = st.text_area(
    "Digite aqui observações que se aplicam a todos os casos:",
    value=st.session_state.consideracoes_gerais,
    height=120,
    key="consideracoes_gerais_area",
)

# ---------------------------------------------------------------------------
# Histórico da sessão
# ---------------------------------------------------------------------------
if st.session_state.relatorios_ia:
    st.markdown("---")
    st.header("Histórico da Sessão")

    casos_ordenados = sorted(st.session_state.relatorios_ia.keys(), key=extrair_numero)
    abas = st.tabs([f"{c}" for c in casos_ordenados])
    for i, nome in enumerate(casos_ordenados):
        with abas[i]:
            if st.session_state.identificacao_exames.get(nome, "").strip():
                st.markdown(f"**Identificação do Exame:** {st.session_state.identificacao_exames[nome]}")
            with st.expander("Ver texto bruto"):
                st.caption(st.session_state.casos_salvos[nome])
            if st.session_state.consideracoes_caso.get(nome, "").strip():
                st.info(st.session_state.consideracoes_caso[nome])
            st.markdown("**Relatório gerado:**")
            st.write(st.session_state.relatorios_ia[nome])

# ---------------------------------------------------------------------------
# Relatório geral (quando há pelo menos 2 casos)
# ---------------------------------------------------------------------------
if len(st.session_state.casos_salvos) >= 2:
    st.markdown("---")
    if st.button("Gerar Relatório Geral", type="primary", use_container_width=True):
        compilado = "".join([f"\n[{k}]: {v}\n" for k, v in st.session_state.casos_salvos.items()])
        texto_geral_para_ia = compilado
        if st.session_state.consideracoes_gerais.strip():
            texto_geral_para_ia += f"\n\nConsiderações gerais do avaliador: {st.session_state.consideracoes_gerais}"

        prompt_geral = (
            "Com base nos relatórios individuais abaixo, elabore um único parágrafo resumindo os achados gerais. "
            "Não mencione os números dos casos, apenas faça um resumo conciso.\n\n"
            f"Relatórios:\n{texto_geral_para_ia}"
        )
        try:
            response_geral = model.generate_content(prompt_geral)
            st.session_state.relatorio_geral_salvo = response_geral.text
            st.session_state.docx_bytes = None
            st.success("Relatório geral gerado!")
            st.rerun()
        except Exception as e:
            st.error(f"Erro ao gerar relatório geral: {e}")

    if st.session_state.relatorio_geral_salvo:
        st.markdown("### Relatório Geral ")
        st.info(st.session_state.relatorio_geral_salvo)
        st.markdown("---")
        casos_ordenados = sorted(st.session_state.relatorios_ia.keys(), key=extrair_numero)
        for nome_grupo, questoes in perguntas.items():
            st.subheader(f"Tabela de Respostas - {nome_grupo}")
            perguntas_ordenadas = list(questoes.keys())
            st.markdown(gerar_tabela_html(casos_ordenados, perguntas_ordenadas, st.session_state.escolhas_casos), unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Exportação do documento Word
# ---------------------------------------------------------------------------
if st.session_state.relatorios_ia:
    st.markdown("---")
    st.header("Exportar Documento")

    def limpar_formatacao(texto):
        return texto.replace("**", "").replace("__", "").replace("#", "")

    with st.expander("Visualizar Prévia do Documento", expanded=True):
        cab = st.session_state.dados_cabecalho
        st.markdown(f"""
        <div class='preview-box'>
            <strong>Cabeçalho</strong><br>
            <strong>Mamógrafo:</strong> {cab['mamografo_fabricante']} - {cab['mamografo_modelo']}<br>
            <strong>CNES:</strong> {cab['cnes']} &nbsp;|&nbsp; <strong>QIID:</strong> {cab['qiid']}<br>
            <strong>Tipo:</strong> {cab['tipo_mamografo']}<br>
            <strong>Instituição:</strong> {cab['instituicao']}<br>
            <strong>Cidade/Estado:</strong> {cab['cidade']} - {cab['estado']}
        </div>
        """, unsafe_allow_html=True)

        casos_ordenados = sorted(st.session_state.relatorios_ia.keys(), key=extrair_numero)
        for nome_grupo, questoes in perguntas.items():
            st.markdown(f"**Tabela de Respostas - {nome_grupo}**")
            perguntas_ordenadas = list(questoes.keys())
            st.markdown(gerar_tabela_html(casos_ordenados, perguntas_ordenadas, st.session_state.escolhas_casos), unsafe_allow_html=True)

        st.markdown("**Identificação dos Exames**")
        id_tabela = "| Caso | Identificação do Exame |\n| --- | --- |\n"
        for caso in casos_ordenados:
            id_texto = st.session_state.identificacao_exames.get(caso, "")
            id_tabela += f"| {caso} | {id_texto} |\n"
        st.markdown(id_tabela)
        st.markdown("---")

        for nome_caso in casos_ordenados:
            st.markdown(f"**{nome_caso}**")
            st.write(st.session_state.relatorios_ia[nome_caso])
            if st.session_state.consideracoes_caso.get(nome_caso, "").strip():
                st.info(st.session_state.consideracoes_caso[nome_caso])
            st.markdown("---")
        if st.session_state.relatorio_geral_salvo:
            st.markdown("**Relatório Geral **")
            st.write(st.session_state.relatorio_geral_salvo)

    def criar_docx_limpo():
        doc = Document()

        doc.add_heading("Instrumento para a análise da qualidade da mamografia", level=0)
        doc.add_paragraph()

        cab = st.session_state.dados_cabecalho
        p = doc.add_paragraph()
        p.add_run("Mamógrafo (fabricante e modelo): ").bold = True
        p.add_run(f"{cab['mamografo_fabricante']} - {cab['mamografo_modelo']}")

        p = doc.add_paragraph()
        p.add_run("CNES: ").bold = True
        p.add_run(cab["cnes"])
        p.add_run("     QIID: ").bold = True
        p.add_run(cab["qiid"])

        p = doc.add_paragraph()
        p.add_run("Tipo de mamógrafo: ").bold = True
        opcoes_tipo = ["Convencional", "Digital CR", "Digital DR", "DR retrofit"]
        for opcao in opcoes_tipo:
            marcado = "X" if cab["tipo_mamografo"] == opcao else "-"
            p.add_run(f"  [{marcado}] {opcao}  ")

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Instituição: ").bold = True
        p.add_run(cab["instituicao"])

        p = doc.add_paragraph()
        p.add_run("Cidade: ").bold = True
        p.add_run(cab["cidade"])
        p.add_run("     Estado: ").bold = True
        p.add_run(cab["estado"])
        doc.add_paragraph()

        # Tabelas de respostas - uma por grupo
        casos_ord = sorted(st.session_state.relatorios_ia.keys(), key=extrair_numero)
        num_casos = len(casos_ord)

        for nome_grupo, questoes in perguntas.items():
            doc.add_heading(f"Tabela de Respostas - {nome_grupo}", level=1)
            perguntas_ord = list(questoes.keys())
            total_colunas = 1 + num_casos * 2
            tabela = doc.add_table(rows=2 + len(perguntas_ord), cols=total_colunas)
            tabela.style = "Table Grid"

            tabela.cell(0, 0).merge(tabela.cell(1, 0))
            tabela.cell(0, 0).text = "Pergunta"
            set_cell_shading(tabela.cell(0, 0), "E0E0E0")
            set_cell_shading(tabela.cell(1, 0), "E0E0E0")

            for idx, caso in enumerate(casos_ord):
                col_inicio = 1 + idx * 2
                col_fim = col_inicio + 1
                tabela.cell(0, col_inicio).merge(tabela.cell(0, col_fim))
                cell_caso = tabela.cell(0, col_inicio)
                cell_caso.text = ""
                run = cell_caso.paragraphs[0].add_run(caso)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                set_cell_shading(cell_caso, "E0E0E0")

            for idx in range(num_casos):
                col_sim = 1 + idx * 2
                col_nao = col_sim + 1
                tabela.cell(1, col_sim).text = "Sim"
                set_cell_shading(tabela.cell(1, col_sim), "E0E0E0")
                tabela.cell(1, col_nao).text = "Não"
                set_cell_shading(tabela.cell(1, col_nao), "E0E0E0")

            for i, pergunta in enumerate(perguntas_ord):
                linha_atual = i + 2
                tabela.cell(linha_atual, 0).text = pergunta
                for j, caso in enumerate(casos_ord):
                    item = st.session_state.escolhas_casos.get(caso, {}).get(pergunta, {})
                    resposta = item.get("resposta", "-") if isinstance(item, dict) else item
                    col_sim = 1 + j * 2
                    col_nao = col_sim + 1
                    tabela.cell(linha_atual, col_sim).text = "X" if resposta == "Sim" else ""
                    tabela.cell(linha_atual, col_nao).text = "X" if resposta == "Não" else ""

            doc.add_paragraph()

        doc.add_page_break()

        # Identificação dos exames
        doc.add_heading("Identificação dos Exames", level=2)
        mini_tabela = doc.add_table(rows=2, cols=num_casos)
        mini_tabela.style = "Table Grid"
        for j, caso in enumerate(casos_ord):
            mini_tabela.rows[0].cells[j].text = caso
        for j, caso in enumerate(casos_ord):
            mini_tabela.rows[1].cells[j].text = st.session_state.identificacao_exames.get(caso, "")
        doc.add_paragraph()

        # Anexo
        doc.add_heading("Anexo ao instrumento para análise da qualidade da mamografia", level=0)
        p = doc.add_paragraph()
        p.add_run("Serviço: ").bold = True
        p.add_run(cab["servico"])
        p = doc.add_paragraph()
        p.add_run("CNES: ").bold = True
        p.add_run(cab["cnes"])
        p = doc.add_paragraph()
        p.add_run("QIID: ").bold = True
        p.add_run(cab["qiid"])
        doc.add_paragraph()

        # Considerações específicas
        doc.add_heading("Considerações Específicas", level=0)
        for nome_caso in casos_ord:
            doc.add_heading(nome_caso, level=1)
            texto_ia = st.session_state.relatorios_ia[nome_caso]
            for linha in texto_ia.strip().split("\n"):
                if linha.strip():
                    doc.add_paragraph(limpar_formatacao(linha))
            if st.session_state.consideracoes_caso.get(nome_caso, "").strip():
                doc.add_heading("Considerações Adicionais", level=2)
                doc.add_paragraph(limpar_formatacao(st.session_state.consideracoes_caso[nome_caso]))
            doc.add_paragraph("-" * 30)

        # Recomendações
        # Estrutura: {pergunta: {sub_opcao: texto}}.
        # Perguntas sem sub_opcoes usam a chave "_default".
        # Perguntas com sub_opcoes podem ter um texto por sub-opção; se uma
        # sub-opção não tiver entrada própria, cai no "_default" da pergunta
        # (se existir) — assim dá pra ir preenchendo aos poucos.
        recomendacoes = {
            "Recomendação correta segundo o BI-RADS": {
                "_default":
                    "Para cada classificação é importante descrever a recomendação apropriada, "
                    "segundo a quinta edição do BI-RADS®, conforme determina a Portaria de Consolidação "
                    "nº 5 GM/MS de 28/09/2017, que no seu anexo XXVIII, estabelece: "
                    "\"o laudo radiográfico deve conter as seguintes informações: "
                    "a) identificação do serviço, da idade do examinado e data do exame; "
                    "b) se exame de rastreamento ou de diagnóstico; "
                    "c) número de filmes ou imagens; "
                    "d) padrão mamário; "
                    "e) achados radiográficos; "
                    "f) classificação BI-RADS®; "
                    "g) recomendação de conduta; e "
                    "h) nome e assinatura do médico interpretador do exame.\"",
            },
            "Contraste adequado": {
                # TODO: revisar/ajustar cada texto — estes são placeholders de base.
                "Contraste alto":
                    "As imagens estão com o contraste aumentado devido à acentuada diferença entre os tons de cinza claros e escuros presentes. Para que o contraste das imagens seja considerado adequado, essa diferença deve ser menos acentuada.",
                "Contraste muito alto":
                    "As imagens deste exame estão com o contraste muito alto e com as regiões correspondentes a tecidos mamários mais densos com os tons de cinza claro saturados (muito claros, quase transparentes), o mesmo ocorrendo nas regiões das axilas nas incidências mediolaterais oblíquas (MLO). Este aspecto das imagens dificulta ou mesmo inviabiliza a identificação de microcalcificações nas regiões de tecidos densos. Portanto, as imagens impressas enviadas para avaliação foram consideradas sem qualidade técnica para a interpretação diagnóstica.",
                "Contraste baixo":
                    "As imagens das quatro incidências estão muito claras e, por conseguinte, com o contraste reduzido devido à pouca diferença entre os tons de cinza claros (regiões de tecido fibroglandular) e de cinza escuros (regiões de tecido subcutâneo e tecido adiposo retromamário) presentes. Para que o contraste das imagens seja considerado adequado, essa diferença deve ser mais acentuada.",
            },
            "Definição de estruturas":{
                "_default":
                    "Com vistas a evitar a perda de definição das imagens (imagens tremidas), é recomendado aos profissionais que realizam os exames atenção à compressão correta das mamas e que sinalizem para as pacientes que não se movimentem e prendam a respiração durante a aquisição das imagens."
            },
            "Imagem sem artefatos (se houver, descrever)":{
                "_default":
                    "É recomendado ao pessoal de manutenção do mamógrafo ajustar o movimento da grade antidifusora de modo que as suas linhas de material radiopaco não sejam registradas nas imagens das pacientes, gerando artefatos. O movimento insuficiente da grade antidifusora do mamógrafo também está gerando um nível de ruído (aspecto granulado) perceptível nas imagens das pacientes. Ver no folder “Critérios de Qualidade da Imagem em Mamografia”, enviado em anexo, as descrições de ruído e artefatos de imagem."
            },
            "Identificação correta do exame":{
                "_default":
                    "É recomendado a(o)s técnica(o)s responsáveis pela impressão dos exames não sobrepor textos de identificação do serviço, da paciente e das técnicas radiográficas sobre áreas das imagens das mamas."
            },
            "Adequada compressão de mama":{
                "_default":
                    "Com vistas a evitar a perda de definição das imagens (imagens tremidas), é recomendado aos profissionais que realizam os exames atenção à compressão correta das mamas e que sinalizem para as pacientes que não se movimentem e prendam a respiração durante a aquisição das imagens."
            },
            "Visibilização completa do parênquima mamário":{
                "_default":
                    "Para o posicionamento adequado as papilas devem estar perfiladas e deve ser incluído todo o tecido fibroglandular nas duas incidências (CC e MLO). A incidência MLO deve conter as pregas inframamárias e o músculo peitoral deve estar na altura ou abaixo das papilas. A incidência CC deve apresentar as papilas equidistantes medial e lateralmente. Deve haver insinuação do músculo peitoral na região central e posterior das mamas na incidência CC e a diferença de parênquima aparente entre esta incidência e a incidência MLO deve ser de no máximo 1,0 cm."
            },
            "Músculo grande peitoral na altura do mamilo ou abaixo - na 0ML":{
                "_default":
                    "Melhorar a tração e a elevação das mamas nas incidências mediolaterais oblíquas (MLO) de modo a incluir nas imagens a prega inframamária e o músculo grande peitoral na altura ou abaixo da papila."
            },
            "Prega inframamária incluída na radiografia - na 0ML":{
                "_default":
                    "Melhorar a tração e a elevação das mamas nas incidências mediolaterais oblíquas (MLO) de modo a incluir nas imagens a prega inframamária e o músculo grande peitoral na altura ou abaixo da papila."
            },
        }

        def resposta_do_caso(caso, pergunta):
            item = st.session_state.escolhas_casos.get(caso, {}).get(pergunta, {})
            return item.get("resposta", "") if isinstance(item, dict) else item

        def sub_opcao_do_caso(caso, pergunta):
            item = st.session_state.escolhas_casos.get(caso, {}).get(pergunta, {})
            return item.get("sub_opcao") if isinstance(item, dict) else None

        def obter_gatilho(pergunta):
            # Busca o gatilho configurado em `perguntas` (padrão: "Não").
            for questoes in perguntas.values():
                if pergunta in questoes:
                    return questoes[pergunta].get("gatilho_sub_opcoes", "Não")
            return "Não"

        tem_recomendacao = any(
            resposta_do_caso(caso, pergunta) == obter_gatilho(pergunta)
            for caso in casos_ord
            for pergunta in recomendacoes
        )
        if tem_recomendacao:
            doc.add_heading("Recomendações", level=0)
            textos_inseridos = set()
            for caso in casos_ord:
                for pergunta, textos_por_sub in recomendacoes.items():
                    if resposta_do_caso(caso, pergunta) == obter_gatilho(pergunta):
                        subs = sub_opcao_do_caso(caso, pergunta)
                        if not subs:
                            # Pergunta sem sub_opcoes (ou nenhuma sub-opção marcada): usa o texto padrão.
                            textos = [textos_por_sub.get("_default")]
                        else:
                            # Uma ou mais sub-opções marcadas: junta o texto de cada uma.
                            textos = [
                                textos_por_sub.get(sub) or textos_por_sub.get("_default")
                                for sub in subs
                            ]
                        for texto in textos:
                            if texto and texto not in textos_inseridos:
                                textos_inseridos.add(texto)
                                p = doc.add_paragraph(style="List Bullet")
                                p.add_run(texto)

        if st.session_state.relatorio_geral_salvo:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Todos os casos:").bold = True
            for linha in st.session_state.relatorio_geral_salvo.strip().split("\n"):
                if linha.strip():
                    doc.add_paragraph(limpar_formatacao(linha))

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    if st.button("Gerar Documento Word", use_container_width=True):
        with st.spinner("Montando o documento..."):
            st.session_state.docx_bytes = criar_docx_limpo().getvalue()
        st.success("Documento gerado! Use o botão abaixo para baixar.")

    if st.session_state.get("docx_bytes"):
        st.download_button(
            label="Baixar Documento Final (.docx)",
            data=st.session_state.docx_bytes,
            file_name="relatorio_final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ---------------------------------------------------------------------------
# Botão de reset
# ---------------------------------------------------------------------------
st.markdown("---")
if st.button("Limpar todos os dados da sessão"):
    for chave in [
        "casos_salvos",
        "relatorios_ia",
        "relatorio_geral_salvo",
        "consideracoes_caso",
        "consideracoes_gerais",
        "escolhas_casos",
        "dados_cabecalho",
        "identificacao_exames",
        "docx_bytes",
    ]:
        if chave in st.session_state:
            del st.session_state[chave]
    st.rerun()
